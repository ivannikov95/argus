from __future__ import annotations

import io
import json
import math
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from PIL import Image, ImageDraw
from pydantic import BaseModel, Field
from scipy import optimize, stats


ROOT = Path(__file__).resolve().parents[1]
PROJECTS_DIR = ROOT / "data" / "projects"
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="Argus API", version="0.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class TableOneRequest(BaseModel):
    rows: list[dict[str, Any]]
    group_column: str | None = None
    variables: list[str] | None = None
    variable_overrides: dict[str, dict[str, Any]] = {}
    numeric_presentation: str = "auto"
    numeric_test: str = "auto"
    categorical_test: str = "auto"
    confidence_level: float = Field(default=0.95, gt=0.5, lt=1.0)


class ProjectSaveRequest(BaseModel):
    project_id: str | None = None  # if set → overwrite existing file
    project_name: str = Field(min_length=1, max_length=120)
    file_name: str = "dataset"
    rows: list[dict[str, Any]]
    variable_overrides: dict[str, dict[str, Any]] = {}
    slides: list[dict[str, Any]] = []          # all table slides
    regression: dict[str, Any] | None = None
    last_analysis: dict[str, Any] | None = None  # backward compat
    table_settings: dict[str, Any] = {}          # backward compat


class ExportRequest(BaseModel):
    title: str = "Таблица 1 - Характеристики выборки"
    description: str = ""
    footnotes: str = ""
    show_overall: bool = True
    show_effect: bool = True
    show_ci: bool = True
    show_missing: bool = True
    decompose_categories: bool = False
    analysis: dict[str, Any]


class DatasetProfileRequest(BaseModel):
    rows: list[dict[str, Any]]
    file_name: str = "Загруженный датасет"


class RegressionRequest(BaseModel):
    rows: list[dict[str, Any]]
    outcome: str
    predictors: list[str] = Field(min_length=1)
    variable_overrides: dict[str, dict[str, Any]] = {}
    confidence_level: float = Field(default=0.95, gt=0.5, lt=1.0)


class CorrelationPairData(BaseModel):
    row: str
    col: str
    x_values: list[float | None]
    y_values: list[float | None]
    x_label: str = ""
    y_label: str = ""


class CorrelationReportSection(BaseModel):
    result: dict[str, Any]
    include_matrix: bool = True
    pairs: list[CorrelationPairData] = []


class ReportExportRequest(BaseModel):
    tables: list[ExportRequest] = []
    regression: dict[str, Any] | None = None
    correlation: CorrelationReportSection | None = None


def clean_scalar(value: Any) -> Any:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return None
    if isinstance(value, (np.integer,)):
        return int(value)
    if isinstance(value, (np.floating,)):
        return float(value)
    if isinstance(value, (pd.Timestamp, datetime)):
        return value.isoformat()
    return value


def frame_to_rows(df: pd.DataFrame) -> list[dict[str, Any]]:
    return [
        {str(key): clean_scalar(value) for key, value in row.items()}
        for row in df.to_dict(orient="records")
    ]


def to_frame(rows: list[dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        raise HTTPException(422, "Датасет пуст")
    return pd.DataFrame(rows)


def infer_column(series: pd.Series) -> dict[str, Any]:
    non_null = series.dropna()
    unique = int(non_null.nunique(dropna=True))
    numeric = pd.to_numeric(non_null, errors="coerce")
    numeric_share = float(numeric.notna().mean()) if len(non_null) else 0.0
    if unique == 2:
        kind = "binary"
    elif numeric_share >= 0.9 and unique > 2:
        kind = "numeric"
    elif unique <= max(20, int(math.sqrt(max(len(series), 1))) + 2):
        kind = "categorical"
    else:
        kind = "text"
    examples = [clean_scalar(v) for v in non_null.unique()[:3]]
    return {
        "name": str(series.name),
        "label": str(series.name).replace("_", " ").strip().title(),
        "type": kind,
        "role": "id" if str(series.name).lower() in {"id", "patient_id", "subject_id"} else "feature",
        "count": int(non_null.size),
        "missing": int(series.isna().sum()),
        "missing_percent": round(float(series.isna().mean() * 100), 1),
        "unique": unique,
        "examples": examples,
    }


def dataset_profile(df: pd.DataFrame) -> dict[str, Any]:
    schema = [infer_column(df[column]) for column in df.columns]
    return {
        "row_count": int(len(df)),
        "column_count": int(len(df.columns)),
        "missing_count": int(df.isna().sum().sum()),
        "duplicate_count": int(df.duplicated().sum()),
        "schema": schema,
        "rows": frame_to_rows(df),
    }


def demo_frame() -> pd.DataFrame:
    rng = np.random.default_rng(20250628)
    n = 120
    group = np.where(np.arange(n) % 2 == 0, "Контроль", "Терапия")
    age = np.clip(rng.normal(62 + (group == "Терапия") * 1.8, 10, n), 31, 86).round()
    bmi = rng.normal(27.5 - (group == "Терапия") * 0.4, 4.1, n).round(1)
    crp = np.exp(rng.normal(1.05 - (group == "Терапия") * 0.2, 0.8, n)).round(1)
    lvef = np.clip(rng.normal(52 + (group == "Терапия") * 2.2, 7, n), 25, 72).round()
    sex = rng.choice(["Женский", "Мужской"], n, p=[0.44, 0.56])
    hypertension = np.where(rng.random(n) < (0.35 + age / 250), "Да", "Нет")
    event = np.where(rng.random(n) < (0.23 - (group == "Терапия") * 0.06), "Да", "Нет")
    df = pd.DataFrame({
        "patient_id": [f"P-{i:03d}" for i in range(1, n + 1)],
        "group": group,
        "age": age,
        "sex": sex,
        "bmi": bmi,
        "crp": crp,
        "lvef": lvef,
        "hypertension": hypertension,
        "event_30d": event,
    })
    df.loc[[7, 41, 88], "crp"] = np.nan
    df.loc[[15, 70], "bmi"] = np.nan
    return df


def welch_anova(groups: list[pd.Series]) -> tuple[float, float]:
    k = len(groups)
    ns = np.array([len(g) for g in groups], dtype=float)
    means = np.array([float(g.mean()) for g in groups])
    vars_ = np.array([float(g.var(ddof=1)) for g in groups])
    if np.any(ns < 2) or np.any(vars_ == 0):
        f_stat, p_value = stats.f_oneway(*groups)
        return float(f_stat), float(p_value)
    ws = ns / vars_
    W = ws.sum()
    grand = (ws * means).sum() / W
    tmp = ((1 - ws / W) ** 2 / (ns - 1)).sum()
    F = (ws * (means - grand) ** 2).sum() / ((k - 1) * (1 + 2 * (k - 2) / (k ** 2 - 1) * tmp))
    df2 = (k ** 2 - 1) / (3 * tmp)
    return float(F), float(stats.f.sf(F, k - 1, df2))


def normal_enough(values: pd.Series) -> bool:
    values = values.dropna().astype(float)
    if len(values) < 8:
        return False
    if len(values) > 5000:
        values = values.sample(5000, random_state=42)
    try:
        return bool(stats.shapiro(values).pvalue >= 0.05)
    except ValueError:
        return False


def fmt_number(value: float, digits: int = 1) -> str:
    if not np.isfinite(value):
        return "—"
    return f"{value:.{digits}f}".replace(".", ",")


def numeric_summary(values: pd.Series, use_mean: bool) -> str:
    values = pd.to_numeric(values, errors="coerce").dropna()
    if values.empty:
        return "—"
    if use_mean:
        return f"{fmt_number(values.mean())} ± {fmt_number(values.std(ddof=1))}"
    q1, median, q3 = values.quantile([0.25, 0.5, 0.75])
    return f"{fmt_number(median)} [{fmt_number(q1)}; {fmt_number(q3)}]"


def categorical_summary(values: pd.Series) -> str:
    clean = values.dropna().astype(str)
    if clean.empty:
        return "—"
    counts = clean.value_counts()
    return "; ".join(
        f"{level}: {count} ({count / len(clean) * 100:.1f}%)"
        for level, count in counts.items()
    )


def p_text(value: float | None) -> str:
    if value is None or not np.isfinite(value):
        return "—"
    return "<0,001" if value < 0.001 else f"{value:.3f}".replace(".", ",")


def mean_difference_ci(first: pd.Series, second: pd.Series, confidence: float) -> tuple[float, float, float] | None:
    if len(first) < 2 or len(second) < 2:
        return None
    estimate = float(second.mean() - first.mean())
    variance_first = float(first.var(ddof=1) / len(first))
    variance_second = float(second.var(ddof=1) / len(second))
    standard_error = math.sqrt(variance_first + variance_second)
    if not standard_error:
        return estimate, estimate, estimate
    denominator = (variance_first**2 / (len(first) - 1)) + (variance_second**2 / (len(second) - 1))
    degrees_freedom = (variance_first + variance_second) ** 2 / denominator if denominator else len(first) + len(second) - 2
    critical = float(stats.t.ppf((1 + confidence) / 2, degrees_freedom))
    return estimate, estimate - critical * standard_error, estimate + critical * standard_error


def median_difference_ci(first: pd.Series, second: pd.Series, confidence: float) -> tuple[float, float, float] | None:
    if len(first) < 2 or len(second) < 2:
        return None
    first_values = first.to_numpy(dtype=float)
    second_values = second.to_numpy(dtype=float)
    rng = np.random.default_rng(20250628)
    if len(first_values) > 5000:
        first_values = rng.choice(first_values, 5000, replace=False)
    if len(second_values) > 5000:
        second_values = rng.choice(second_values, 5000, replace=False)
    estimate = float(np.median(second_values) - np.median(first_values))
    bootstrap = np.empty(1000)
    for index in range(len(bootstrap)):
        bootstrap[index] = np.median(rng.choice(second_values, len(second_values), replace=True)) - np.median(rng.choice(first_values, len(first_values), replace=True))
    alpha = (1 - confidence) / 2
    lower, upper = np.quantile(bootstrap, [alpha, 1 - alpha])
    return estimate, float(lower), float(upper)


def odds_ratio_ci(table: pd.DataFrame, confidence: float) -> tuple[float, float, float, str] | None:
    if table.shape != (2, 2):
        return None
    values = table.to_numpy(dtype=float)
    event_control, event_treatment = values[1, 0], values[1, 1]
    no_event_control, no_event_treatment = values[0, 0], values[0, 1]
    if (values == 0).any():
        event_control += 0.5
        event_treatment += 0.5
        no_event_control += 0.5
        no_event_treatment += 0.5
    odds_ratio = (event_treatment * no_event_control) / (no_event_treatment * event_control)
    standard_error = math.sqrt(1 / event_treatment + 1 / no_event_treatment + 1 / event_control + 1 / no_event_control)
    critical = float(stats.norm.ppf((1 + confidence) / 2))
    lower = math.exp(math.log(odds_ratio) - critical * standard_error)
    upper = math.exp(math.log(odds_ratio) + critical * standard_error)
    return float(odds_ratio), float(lower), float(upper), str(table.index[1])


def level_summaries(series: pd.Series, grouped: list[pd.Series], groups: list[str]) -> list[dict[str, Any]]:
    levels = [str(value) for value in series.dropna().astype(str).unique()]
    levels = sorted(levels)
    result = []
    for level in levels:
        clean = series.dropna().astype(str)
        overall_count = int((clean == level).sum())
        group_values: dict[str, str] = {}
        for group_name, values in zip(groups, grouped):
            group_clean = values.dropna().astype(str)
            count = int((group_clean == level).sum())
            group_values[group_name] = f"{count} ({count / len(group_clean) * 100:.1f}%)" if len(group_clean) else "—"
        result.append({
            "level": level,
            "overall": f"{overall_count} ({overall_count / len(clean) * 100:.1f}%)" if len(clean) else "—",
            "groups": group_values,
        })
    return result


def analyze_numeric(df: pd.DataFrame, variable: str, group: str | None, groups: list[str], request: TableOneRequest) -> dict[str, Any]:
    series = pd.to_numeric(df[variable], errors="coerce")
    grouped = [pd.to_numeric(df.loc[df[group].astype(str) == g, variable], errors="coerce").dropna() for g in groups] if group else []
    automatic_mean = normal_enough(series) and all(normal_enough(part) for part in grouped)
    use_mean = automatic_mean if request.numeric_presentation == "auto" else request.numeric_presentation == "mean_sd"
    parametric_test = use_mean if request.numeric_test == "auto" else request.numeric_test == "parametric"
    test_name, p_value, effect, ci = "—", None, None, None
    ci_label = "—"
    if len(grouped) == 2 and all(len(part) >= 2 for part in grouped):
        if parametric_test:
            result = stats.ttest_ind(grouped[0], grouped[1], equal_var=False, nan_policy="omit")
            test_name, p_value = "t-критерий Уэлча", float(result.pvalue)
            ci = mean_difference_ci(grouped[0], grouped[1], request.confidence_level)
            ci_label = "Разность средних"
        else:
            result = stats.mannwhitneyu(grouped[0], grouped[1], alternative="two-sided")
            test_name, p_value = "U-критерий Манна–Уитни", float(result.pvalue)
            ci = median_difference_ci(grouped[0], grouped[1], request.confidence_level)
            ci_label = "Разность медиан (bootstrap)"
        pooled = math.sqrt((grouped[0].var(ddof=1) + grouped[1].var(ddof=1)) / 2)
        effect = float((grouped[1].mean() - grouped[0].mean()) / pooled) if pooled else None
    elif len(grouped) > 2 and all(len(part) >= 2 for part in grouped):
        if parametric_test:
            _, p_value = welch_anova(grouped)
            test_name = "ANOVA Уэлча"
        else:
            result = stats.kruskal(*grouped)
            test_name, p_value = "Краскела–Уоллиса", float(result.pvalue)
    confidence_percent = round(request.confidence_level * 100)
    return {
        "variable": variable,
        "type": "numeric",
        "presentation": "Среднее ± SD" if use_mean else "Медиана [Q1; Q3]",
        "overall": numeric_summary(series, use_mean),
        "groups": {g: numeric_summary(part, use_mean) for g, part in zip(groups, grouped)},
        "missing": int(series.isna().sum()),
        "test": test_name,
        "p_value": p_value,
        "p_display": p_text(p_value),
        "effect": fmt_number(effect, 2) if effect is not None else "—",
        "effect_label": "SMD" if effect is not None else "—",
        "ci_display": f"{fmt_number(ci[0], 2)} [{fmt_number(ci[1], 2)}; {fmt_number(ci[2], 2)}]" if ci else "—",
        "ci_label": f"{ci_label}, {confidence_percent}% ДИ" if ci else "—",
        "normality": "Распределение совместимо с нормальным" if automatic_mean else "Есть отклонение от нормального распределения",
        "levels": [],
    }


def analyze_categorical(df: pd.DataFrame, variable: str, group: str | None, groups: list[str], request: TableOneRequest) -> dict[str, Any]:
    series = df[variable]
    grouped = [df.loc[df[group].astype(str) == g, variable] for g in groups] if group else []
    test_name, p_value, effect, ci = "—", None, None, None
    ci_label = "—"
    if group and len(groups) >= 2:
        table = pd.crosstab(df[variable], df[group])
        if table.shape[0] >= 2 and table.shape[1] >= 2:
            chi2, chi_p, _, expected = stats.chi2_contingency(table)
            force_fisher = request.categorical_test == "fisher" and table.shape == (2, 2)
            force_chi = request.categorical_test == "chi_square"
            if table.shape == (2, 2) and ((expected < 5).any() or force_fisher) and not force_chi:
                _, fisher_p = stats.fisher_exact(table)
                test_name, p_value = "Точный критерий Фишера", float(fisher_p)
            else:
                test_name, p_value = "χ² Пирсона", float(chi_p)
            denom = min(table.shape[0] - 1, table.shape[1] - 1)
            effect = math.sqrt(chi2 / (table.to_numpy().sum() * denom)) if denom else None
            ci = odds_ratio_ci(table, request.confidence_level)
            if ci:
                ci_label = f"ОШ для уровня «{ci[3]}», {round(request.confidence_level * 100)}% ДИ"
    return {
        "variable": variable,
        "type": "categorical",
        "presentation": "n (%)",
        "overall": categorical_summary(series),
        "groups": {g: categorical_summary(part) for g, part in zip(groups, grouped)},
        "missing": int(series.isna().sum()),
        "test": test_name,
        "p_value": p_value,
        "p_display": p_text(p_value),
        "effect": fmt_number(effect, 2) if effect is not None else "—",
        "effect_label": "V Крамера" if effect is not None else "—",
        "ci_display": f"{fmt_number(ci[0], 2)} [{fmt_number(ci[1], 2)}; {fmt_number(ci[2], 2)}]" if ci else "—",
        "ci_label": ci_label,
        "normality": "—",
        "levels": level_summaries(series, grouped, groups),
    }


def run_table_one(request: TableOneRequest) -> dict[str, Any]:
    df = to_frame(request.rows)
    if request.group_column and request.group_column not in df.columns:
        raise HTTPException(422, "Группирующая переменная отсутствует в датасете")
    groups = []
    if request.group_column:
        def _fmt_group(v: Any) -> str:
            if isinstance(v, float) and not math.isnan(v) and v == int(v):
                return str(int(v))
            return str(v)
        groups = [_fmt_group(v) for v in df[request.group_column].dropna().unique()]
        groups = sorted(groups)[:8]
        # Normalise the group column in-place so comparisons work correctly
        df = df.copy()
        df[request.group_column] = df[request.group_column].map(_fmt_group)
    variables = request.variables or list(df.columns)
    variables = [v for v in variables if v in df.columns and v != request.group_column]
    schema = {item["name"]: item for item in (infer_column(df[col]) for col in df.columns)}
    # Apply user overrides (type/role changes from the Variables page)
    for col, override in request.variable_overrides.items():
        if col in schema:
            if "type" in override:
                schema[col]["type"] = override["type"]
            if "role" in override:
                schema[col]["role"] = override["role"]
    result_rows = []
    for variable in variables:
        kind = schema[variable]["type"]
        if schema[variable]["role"] == "id" or kind == "text":
            continue
        if kind == "numeric":
            result_rows.append(analyze_numeric(df, variable, request.group_column, groups, request))
        else:
            result_rows.append(analyze_categorical(df, variable, request.group_column, groups, request))
    automatic_mode = request.numeric_presentation == "auto" and request.numeric_test == "auto" and request.categorical_test == "auto"
    method_note = (
        "Представление и критерии выбраны автоматически по типу и распределению данных."
        if automatic_mode
        else "Часть представлений или критериев задана исследователем вручную."
    )
    return {
        "n": int(len(df)),
        "group_column": request.group_column,
        "groups": [{"name": g, "n": int((df[request.group_column].astype(str) == g).sum())} for g in groups] if request.group_column else [],
        "rows": result_rows,
        "note": f"{method_note} Перед публикацией метод должен быть подтверждён исследователем.",
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


def _positive_level(levels: list[str]) -> str:
    preferred = {"1", "да", "yes", "true", "event", "событие", "умер", "death", "positive"}
    return next((level for level in levels if level.strip().lower() in preferred), levels[-1])


def _regression_design(request: RegressionRequest) -> tuple[pd.DataFrame, np.ndarray, np.ndarray, list[str], list[str], str, str | None]:
    df = to_frame(request.rows)
    if request.outcome not in df.columns:
        raise HTTPException(422, "Исход отсутствует в датасете")
    predictors = list(dict.fromkeys(request.predictors))
    if request.outcome in predictors:
        raise HTTPException(422, "Исход нельзя одновременно использовать как предиктор")
    missing = [name for name in predictors if name not in df.columns]
    if missing:
        raise HTTPException(422, f"Предикторы отсутствуют в датасете: {', '.join(missing)}")

    schema = {column: infer_column(df[column]) for column in [request.outcome, *predictors]}
    for column, override in request.variable_overrides.items():
        if column in schema:
            schema[column].update({key: override[key] for key in ("type", "role") if key in override})

    work = df[[request.outcome, *predictors]].copy().dropna()
    numeric_columns = [name for name in [request.outcome, *predictors] if schema[name]["type"] == "numeric"]
    for name in numeric_columns:
        work[name] = pd.to_numeric(work[name], errors="coerce")
    work = work.dropna()
    if work.empty:
        raise HTTPException(422, "После исключения пропусков не осталось наблюдений")

    outcome_levels = sorted(work[request.outcome].astype(str).unique().tolist())
    logistic = schema[request.outcome]["type"] == "binary" or len(outcome_levels) == 2
    event_level: str | None = None
    if logistic:
        if len(outcome_levels) != 2:
            raise HTTPException(422, "Для логистической регрессии исход должен иметь ровно два уровня")
        event_level = _positive_level(outcome_levels)
        y = (work[request.outcome].astype(str) == event_level).astype(float).to_numpy()
    else:
        y_series = pd.to_numeric(work[request.outcome], errors="coerce")
        valid = y_series.notna()
        work, y = work.loc[valid], y_series.loc[valid].to_numpy(dtype=float)
        if len(np.unique(y)) < 2:
            raise HTTPException(422, "Числовой исход не варьирует")

    columns: list[np.ndarray] = [np.ones(len(work), dtype=float)]
    terms = ["Константа"]
    references: list[str] = []
    for predictor in predictors:
        if schema[predictor]["type"] == "numeric":
            values = pd.to_numeric(work[predictor], errors="coerce").to_numpy(dtype=float)
            if np.nanstd(values) == 0:
                continue
            columns.append(values)
            terms.append(predictor)
            continue
        levels = sorted(work[predictor].astype(str).unique().tolist())
        if len(levels) < 2:
            continue
        reference = levels[0]
        references.append(f"{predictor}: {reference}")
        values = work[predictor].astype(str)
        for level in levels[1:]:
            columns.append((values == level).astype(float).to_numpy())
            terms.append(f"{predictor}: {level} (реф. {reference})")

    if len(columns) == 1:
        raise HTTPException(422, "Выбранные предикторы не варьируют")
    x = np.column_stack(columns)
    if len(work) <= x.shape[1] + 1:
        raise HTTPException(422, "Недостаточно полных наблюдений для числа параметров модели")
    if np.linalg.matrix_rank(x) < x.shape[1]:
        raise HTTPException(422, "Предикторы линейно зависимы; уберите дублирующие показатели")
    return work, y, x, terms, references, "logistic" if logistic else "linear", event_level


def _logistic_diagnostics(y: np.ndarray, probabilities: np.ndarray) -> dict[str, Any]:
    positives = int(y.sum())
    negatives = int(len(y) - positives)
    ranks = stats.rankdata(probabilities)
    auc = float((ranks[y == 1].sum() - positives * (positives + 1) / 2) / (positives * negatives))

    order = np.argsort(-probabilities, kind="stable")
    sorted_probabilities = probabilities[order]
    sorted_y = y[order]
    true_positives = np.cumsum(sorted_y)
    false_positives = np.cumsum(1 - sorted_y)
    boundaries = np.flatnonzero(np.r_[sorted_probabilities[:-1] != sorted_probabilities[1:], True])
    roc_curve = [{"fpr": 0.0, "tpr": 0.0, "threshold": 1.0}]
    roc_curve.extend({
        "fpr": float(false_positives[index] / negatives),
        "tpr": float(true_positives[index] / positives),
        "threshold": float(sorted_probabilities[index]),
    } for index in boundaries)
    if len(roc_curve) > 260:
        keep = np.unique(np.linspace(0, len(roc_curve) - 1, 260, dtype=int))
        roc_curve = [roc_curve[index] for index in keep]
    return {
        "auc": auc,
        "roc_curve": roc_curve,
        "predictions": [
            {"actual": int(actual), "probability": float(probability)}
            for actual, probability in zip(y, probabilities)
        ],
    }


def _fit_firth_logistic(x: np.ndarray, y: np.ndarray) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    beta = np.zeros(x.shape[1], dtype=float)

    def penalized_log_likelihood(candidate: np.ndarray) -> float:
        eta = np.clip(x @ candidate, -40, 40)
        probabilities = 1 / (1 + np.exp(-eta))
        weights = np.maximum(probabilities * (1 - probabilities), 1e-10)
        information = x.T @ (x * weights[:, None])
        sign, log_determinant = np.linalg.slogdet(information)
        if sign <= 0:
            return -math.inf
        return float(y @ eta - np.logaddexp(0, eta).sum() + 0.5 * log_determinant)

    for _ in range(150):
        eta = np.clip(x @ beta, -40, 40)
        probabilities = 1 / (1 + np.exp(-eta))
        weights = np.maximum(probabilities * (1 - probabilities), 1e-10)
        information = x.T @ (x * weights[:, None])
        try:
            covariance = np.linalg.inv(information)
        except np.linalg.LinAlgError as exc:
            raise HTTPException(422, "Firth-модель не может быть оценена для выбранных предикторов") from exc
        leverage = weights * np.einsum("ij,jk,ik->i", x, covariance, x)
        adjusted_score = x.T @ (y - probabilities + leverage * (0.5 - probabilities))
        step = covariance @ adjusted_score
        if np.max(np.abs(step)) < 1e-8:
            return beta, covariance, probabilities
        current_likelihood = penalized_log_likelihood(beta)
        scale = 1.0
        while scale > 1e-6 and penalized_log_likelihood(beta + scale * step) < current_likelihood:
            scale *= 0.5
        beta = beta + scale * step
    raise HTTPException(422, "Firth-модель не сошлась; сократите число предикторов")


def run_regression(request: RegressionRequest) -> dict[str, Any]:
    work, y, x, terms, references, model_type, event_level = _regression_design(request)
    confidence = request.confidence_level
    alpha = 1 - confidence
    diagnostics = None
    warnings: list[str] = []
    fit_method = "ols" if model_type == "linear" else "mle"

    if model_type == "linear":
        beta, _, _, _ = np.linalg.lstsq(x, y, rcond=None)
        fitted = x @ beta
        residuals = y - fitted
        degrees_freedom = len(y) - x.shape[1]
        sigma2 = float(residuals @ residuals / degrees_freedom)
        covariance = sigma2 * np.linalg.inv(x.T @ x)
        standard_errors = np.sqrt(np.diag(covariance))
        statistics = beta / standard_errors
        p_values = 2 * stats.t.sf(np.abs(statistics), degrees_freedom)
        critical = float(stats.t.ppf(1 - alpha / 2, degrees_freedom))
        lower, upper = beta - critical * standard_errors, beta + critical * standard_errors
        total_sum = float(((y - y.mean()) ** 2).sum())
        residual_sum = float((residuals**2).sum())
        r_squared = 1 - residual_sum / total_sum if total_sum else 0.0
        adjusted = 1 - (1 - r_squared) * (len(y) - 1) / degrees_freedom
        metrics = {"r_squared": r_squared, "adjusted_r_squared": adjusted, "rmse": math.sqrt(residual_sum / len(y))}
    else:
        def objective(beta: np.ndarray) -> float:
            eta = x @ beta
            return float(np.logaddexp(0, eta).sum() - y @ eta)

        def gradient(beta: np.ndarray) -> np.ndarray:
            eta = np.clip(x @ beta, -40, 40)
            probabilities = 1 / (1 + np.exp(-eta))
            return x.T @ (probabilities - y)

        fit = optimize.minimize(objective, np.zeros(x.shape[1]), jac=gradient, method="BFGS")
        beta = fit.x
        try:
            eta = np.clip(x @ beta, -40, 40)
            probabilities = 1 / (1 + np.exp(-eta))
            weights = np.maximum(probabilities * (1 - probabilities), 1e-12)
            information = x.T @ (x * weights[:, None])
            covariance = np.linalg.inv(information)
            mle_standard_errors = np.sqrt(np.diag(covariance))
            unstable = (
                (not fit.success and np.linalg.norm(fit.jac) > 1e-4)
                or not np.all(np.isfinite(mle_standard_errors))
                or np.max(np.abs(beta)) > 10
                or np.max(mle_standard_errors) > 10
            )
        except np.linalg.LinAlgError:
            unstable = True
        if unstable:
            beta, covariance, probabilities = _fit_firth_logistic(x, y)
            fit_method = "firth"
            warnings.append(
                "Обнаружено полное или квазиполное разделение исходов. Применена логистическая регрессия Фирта; оценки конечны, но при малом числе событий требуют осторожной интерпретации."
            )
        standard_errors = np.sqrt(np.diag(covariance))
        statistics = beta / standard_errors
        p_values = 2 * stats.norm.sf(np.abs(statistics))
        critical = float(stats.norm.ppf(1 - alpha / 2))
        lower, upper = beta - critical * standard_errors, beta + critical * standard_errors
        log_likelihood = -objective(beta)
        event_rate = float(y.mean())
        null_probability = min(max(event_rate, 1e-12), 1 - 1e-12)
        null_ll = float((y * math.log(null_probability) + (1 - y) * math.log(1 - null_probability)).sum())
        metrics = {
            "pseudo_r_squared": 1 - log_likelihood / null_ll,
            "accuracy": float(((probabilities >= 0.5) == y).mean()),
            "aic": 2 * x.shape[1] - 2 * log_likelihood,
        }
        diagnostics = _logistic_diagnostics(y, probabilities)

    coefficients = []
    for index, term in enumerate(terms):
        effect = math.exp(float(np.clip(beta[index], -700, 700))) if model_type == "logistic" else float(beta[index])
        effect_lower = math.exp(float(np.clip(lower[index], -700, 700))) if model_type == "logistic" else float(lower[index])
        effect_upper = math.exp(float(np.clip(upper[index], -700, 700))) if model_type == "logistic" else float(upper[index])
        coefficients.append({
            "term": term,
            "estimate": float(beta[index]),
            "standard_error": float(standard_errors[index]),
            "p_value": float(p_values[index]),
            "p_display": p_text(float(p_values[index])),
            "ci_lower": float(lower[index]),
            "ci_upper": float(upper[index]),
            "effect": effect,
            "effect_ci_lower": effect_lower,
            "effect_ci_upper": effect_upper,
        })
    return {
        "model_type": model_type,
        "outcome": request.outcome,
        "event_level": event_level,
        "n": int(len(work)),
        "excluded": int(len(request.rows) - len(work)),
        "confidence_level": confidence,
        "references": references,
        "coefficients": coefficients,
        "metrics": metrics,
        "diagnostics": diagnostics,
        "fit_method": fit_method,
        "warnings": warnings,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


class LogisticTableRequest(BaseModel):
    rows: list[dict[str, Any]]
    outcome: str
    predictors: list[str] = Field(min_length=1)
    confidence_level: float = Field(default=0.95, gt=0.5, lt=1.0)
    variable_overrides: dict[str, dict[str, Any]] = {}


def _logistic_one(rows: list[dict], outcome: str, predictor: str, confidence_level: float, overrides: dict) -> dict[str, Any]:
    req = RegressionRequest(rows=rows, outcome=outcome, predictors=[predictor],
                            confidence_level=confidence_level, variable_overrides=overrides)
    result = run_regression(req)
    coeff = next((c for c in result["coefficients"] if c["term"] != "Константа"), None)
    if coeff is None:
        raise HTTPException(422, f"Не удалось оценить коэффициент для {predictor}")
    return {
        "or": coeff["effect"],
        "ci_lower": coeff["effect_ci_lower"],
        "ci_upper": coeff["effect_ci_upper"],
        "p_value": coeff["p_value"],
        "p_display": coeff["p_display"],
        "n": result["n"],
        "n_events": int(sum(1 for r in rows if str(r.get(outcome, "")) == result.get("event_level", ""))),
    }


def _nagelkerke_r2(log_lik: float, null_ll: float, n: int) -> float:
    cox_snell = 1 - math.exp(2 * (null_ll - log_lik) / n)
    max_r2 = 1 - math.exp(2 * null_ll / n)
    return cox_snell / max_r2 if max_r2 else 0.0


@app.post("/api/analyze/logistic-univariate")
def logistic_univariate(request: LogisticTableRequest) -> dict[str, Any]:
    results: dict[str, Any] = {}
    for predictor in request.predictors:
        try:
            results[predictor] = _logistic_one(
                request.rows, request.outcome, predictor,
                request.confidence_level, request.variable_overrides
            )
        except HTTPException as exc:
            results[predictor] = {"error": exc.detail}
    return {"univariate": results}


@app.post("/api/analyze/logistic-multivariate")
def logistic_multivariate(request: LogisticTableRequest) -> dict[str, Any]:
    req = RegressionRequest(rows=request.rows, outcome=request.outcome, predictors=request.predictors,
                            confidence_level=request.confidence_level, variable_overrides=request.variable_overrides)
    result = run_regression(req)

    # Per-predictor coefficients (skip intercept)
    coeffs: dict[str, Any] = {}
    for coeff in result["coefficients"]:
        if coeff["term"] == "Константа":
            continue
        # term may be "predictor" or "predictor: level (реф. x)"
        term = coeff["term"]
        pred = next((p for p in request.predictors if term == p or term.startswith(p + ":")), term)
        coeffs[pred] = {
            "term": term,
            "or": coeff["effect"],
            "ci_lower": coeff["effect_ci_lower"],
            "ci_upper": coeff["effect_ci_upper"],
            "p_value": coeff["p_value"],
            "p_display": coeff["p_display"],
        }

    # Model-level stats
    n = result["n"]
    metrics = result["metrics"]
    df_work = to_frame(request.rows)[[request.outcome, *request.predictors]].dropna()
    y_arr = (df_work[request.outcome].astype(str) == str(result.get("event_level", ""))).astype(float).to_numpy()
    event_rate = float(y_arr.mean())
    null_p = min(max(event_rate, 1e-12), 1 - 1e-12)
    null_ll = float(len(y_arr) * (event_rate * math.log(null_p) + (1 - event_rate) * math.log(1 - null_p)))
    pseudo_r2 = metrics.get("pseudo_r_squared", 0.0)
    log_lik = null_ll * (1 - pseudo_r2) if pseudo_r2 != 1 else 0.0
    lr_chi2 = 2 * (log_lik - null_ll)
    chi2_p = float(stats.chi2.sf(abs(lr_chi2), df=len(request.predictors)))
    nagelkerke = _nagelkerke_r2(log_lik, null_ll, n)

    return {
        "coefficients": coeffs,
        "n": n,
        "n_events": int(y_arr.sum()),
        "chi2": abs(lr_chi2),
        "chi2_p": chi2_p,
        "chi2_p_display": p_text(chi2_p),
        "nagelkerke_r2": nagelkerke,
        "fit_method": result["fit_method"],
        "warnings": result["warnings"],
    }


class ModelingRequest(BaseModel):
    rows: list[dict[str, Any]]
    outcome: str
    predictors: list[str] = Field(min_length=1)
    variable_overrides: dict[str, dict[str, Any]] = {}
    train_size: float = Field(default=0.8, gt=0.5, lt=1.0)
    validation_size: float = Field(default=0.0, ge=0.0, lt=0.5)
    random_seed: int = 42
    tuning_method: str = "none"   # "none" | "grid" | "random"
    cv_folds: int = Field(default=5, ge=2, le=20)
    n_iter: int = Field(default=20, ge=5, le=100)
    confidence_level: float = Field(default=0.95, gt=0.5, lt=1.0)
    cutoff: float = Field(default=0.5, gt=0.0, lt=1.0)


def _split_metrics(y_true, y_prob, cutoff: float, confidence_level: float) -> dict[str, Any]:
    from sklearn.metrics import roc_auc_score, roc_curve as sk_roc
    import numpy as np

    n = len(y_true)
    auc = float(roc_auc_score(y_true, y_prob)) if len(set(y_true)) > 1 else 0.5
    fpr_arr, tpr_arr, thr_arr = sk_roc(y_true, y_prob)
    # Clamp thresholds to [0,1] (sklearn prepends max+eps as sentinel)
    roc_points = [
        {"fpr": float(f), "tpr": float(t), "threshold": float(min(max(th, 0.0), 1.0))}
        for f, t, th in zip(fpr_arr, tpr_arr, thr_arr)
    ]
    # Youden's J = sensitivity + specificity - 1 (maximise)
    youden_j = tpr_arr + (1.0 - fpr_arr) - 1.0
    best_j = int(np.argmax(youden_j))
    youden_cutoff = float(min(max(thr_arr[best_j], 0.0), 1.0))

    y_pred = (y_prob >= cutoff).astype(int)
    tp = int(((y_true == 1) & (y_pred == 1)).sum())
    tn = int(((y_true == 0) & (y_pred == 0)).sum())
    fp = int(((y_true == 0) & (y_pred == 1)).sum())
    fn = int(((y_true == 1) & (y_pred == 0)).sum())

    def ratio(a: int, b: int) -> float:
        return a / b if b else 0.0

    sensitivity = ratio(tp, tp + fn)
    specificity = ratio(tn, tn + fp)
    ppv = ratio(tp, tp + fp)
    npv = ratio(tn, tn + fn)
    efficiency = ratio(tp + tn, n)

    # 95% CI for AUC via DeLong's normal approximation
    q1 = auc / (2 - auc)
    q2 = 2 * auc ** 2 / (1 + auc)
    n_pos = int(y_true.sum())
    n_neg = n - n_pos
    var_auc = (auc * (1 - auc) + (n_pos - 1) * (q1 - auc ** 2) + (n_neg - 1) * (q2 - auc ** 2)) / (n_pos * n_neg)
    z = stats.norm.ppf(1 - (1 - confidence_level) / 2)
    auc_lo = max(0.0, auc - z * math.sqrt(max(var_auc, 0)))
    auc_hi = min(1.0, auc + z * math.sqrt(max(var_auc, 0)))

    return {
        "n": n,
        "n_events": int(y_true.sum()),
        "auc": auc,
        "auc_ci_lower": auc_lo,
        "auc_ci_upper": auc_hi,
        "sensitivity": sensitivity,
        "specificity": specificity,
        "ppv": ppv,
        "npv": npv,
        "efficiency": efficiency,
        "tp": tp, "tn": tn, "fp": fp, "fn": fn,
        "roc_curve": roc_points,
        "youden_cutoff": youden_cutoff,
    }


@app.post("/api/analyze/modeling")
def analyze_modeling(request: ModelingRequest) -> dict[str, Any]:
    from sklearn.linear_model import LogisticRegression
    from sklearn.preprocessing import StandardScaler
    from sklearn.pipeline import Pipeline
    from sklearn.model_selection import train_test_split, GridSearchCV, RandomizedSearchCV
    import numpy as np

    # ── prepare dataframe ────────────────────────────────────────────────────
    df = to_frame(request.rows)
    all_cols = [request.outcome, *request.predictors]
    df = df[[c for c in all_cols if c in df.columns]].copy()

    # apply variable overrides (type coercion)
    for col, ov in request.variable_overrides.items():
        if col not in df.columns:
            continue
        typ = ov.get("type", "")
        if typ == "numeric":
            df[col] = pd.to_numeric(df[col], errors="coerce")
        elif typ in ("binary", "categorical"):
            df[col] = df[col].astype(str)

    df = df.dropna()
    if len(df) < 20:
        raise HTTPException(422, "Слишком мало наблюдений после удаления пропусков.")

    # encode outcome
    outcome_vals = sorted(df[request.outcome].astype(str).unique())
    if len(outcome_vals) != 2:
        raise HTTPException(422, f"Исход должен быть бинарным (найдено {len(outcome_vals)} уровня).")
    event_level = outcome_vals[1]  # higher value = event (e.g. "1")
    y_all = (df[request.outcome].astype(str) == event_level).astype(int).to_numpy()

    # encode predictors (dummy-encode categoricals)
    X_raw = df[request.predictors]
    num_cols = X_raw.select_dtypes(include="number").columns.tolist()
    cat_cols = [c for c in request.predictors if c not in num_cols]
    X_enc = pd.get_dummies(X_raw, columns=cat_cols, drop_first=True)
    feature_names = X_enc.columns.tolist()
    X_all = X_enc.to_numpy(dtype=float)

    warnings: list[str] = []
    if len(feature_names) > len(df) // 10:
        warnings.append("Много предикторов относительно объёма выборки — риск переобучения.")

    # ── train / test / validation split ─────────────────────────────────────
    val_size = request.validation_size
    test_size = round(1.0 - request.train_size - val_size, 10)
    if test_size <= 0:
        raise HTTPException(422, "Сумма train + validation не должна превышать 1.")

    if val_size > 0:
        X_tv, X_val, y_tv, y_val = train_test_split(
            X_all, y_all, test_size=val_size, random_state=request.random_seed, stratify=y_all
        )
    else:
        X_tv, y_tv = X_all, y_all
        X_val, y_val = None, None

    X_train, X_test, y_train, y_test = train_test_split(
        X_tv, y_tv,
        test_size=round(test_size / (request.train_size + test_size), 10),
        random_state=request.random_seed,
        stratify=y_tv,
    )

    # ── hyperparameter tuning ────────────────────────────────────────────────
    pipe = Pipeline([
        ("scaler", StandardScaler()),
        ("clf", LogisticRegression(max_iter=1000, random_state=request.random_seed)),
    ])
    best_params: dict[str, Any] | None = None

    if request.tuning_method in ("grid", "random"):
        param_grid = {
            "clf__C": [0.001, 0.01, 0.1, 1, 10, 100],
            "clf__penalty": ["l1", "l2"],
            "clf__solver": ["liblinear"],
        }
        scorer = "roc_auc"
        cv = min(request.cv_folds, int(min(y_train.sum(), (1 - y_train).sum())))
        cv = max(cv, 2)
        if request.tuning_method == "grid":
            search = GridSearchCV(pipe, param_grid, scoring=scorer, cv=cv, n_jobs=-1)
        else:
            search = RandomizedSearchCV(pipe, param_grid, n_iter=min(request.n_iter, 12),
                                        scoring=scorer, cv=cv, random_state=request.random_seed, n_jobs=-1)
        search.fit(X_train, y_train)
        pipe = search.best_estimator_
        best_params = {k.replace("clf__", ""): v for k, v in search.best_params_.items()}
    else:
        pipe.fit(X_train, y_train)

    # ── compute metrics per split ────────────────────────────────────────────
    cutoff = request.cutoff
    train_prob = pipe.predict_proba(X_train)[:, 1]
    test_prob  = pipe.predict_proba(X_test)[:, 1]

    train_metrics = _split_metrics(y_train, train_prob, cutoff, request.confidence_level)
    test_metrics  = _split_metrics(y_test,  test_prob,  cutoff, request.confidence_level)
    val_metrics: dict[str, Any] | None = None
    if X_val is not None:
        val_prob = pipe.predict_proba(X_val)[:, 1]
        val_metrics = _split_metrics(y_val, val_prob, cutoff, request.confidence_level)

    # ── extract coefficients + CIs (from the scaled model) ──────────────────
    clf = pipe.named_steps["clf"]
    scaler = pipe.named_steps["scaler"]
    coefs_raw = clf.coef_[0]
    intercept = float(clf.intercept_[0])

    # Bootstrap CIs for OR
    n_boot = 200
    boot_coefs = np.zeros((n_boot, len(feature_names)))
    rng = np.random.default_rng(request.random_seed)
    for b in range(n_boot):
        idx = rng.integers(0, len(X_train), size=len(X_train))
        try:
            boot_clf = LogisticRegression(
                C=clf.C, penalty=clf.penalty, solver=clf.solver,
                max_iter=500, random_state=request.random_seed
            )
            boot_scaler = StandardScaler()
            Xb = boot_scaler.fit_transform(X_train[idx])
            boot_clf.fit(Xb, y_train[idx])
            boot_coefs[b] = boot_clf.coef_[0]
        except Exception:
            boot_coefs[b] = coefs_raw

    alpha = 1 - request.confidence_level
    coefficients = []
    for i, name in enumerate(feature_names):
        coef = float(coefs_raw[i])
        or_val = math.exp(coef)
        boot_or = np.exp(boot_coefs[:, i])
        ci_lo = float(np.percentile(boot_or, 100 * alpha / 2))
        ci_hi = float(np.percentile(boot_or, 100 * (1 - alpha / 2)))
        p_approx = float(2 * stats.norm.sf(abs(coef / (boot_coefs[:, i].std() + 1e-12))))
        coefficients.append({
            "term": name,
            "coef": coef,
            "or": or_val,
            "ci_lower": ci_lo,
            "ci_upper": ci_hi,
            "p_value": p_approx,
            "p_display": p_text(p_approx),
        })

    return {
        "n_total": len(X_all),
        "n_train": len(X_train),
        "n_test":  len(X_test),
        "n_validation": len(X_val) if X_val is not None else None,
        "event_level": event_level,
        "train": train_metrics,
        "test":  test_metrics,
        "validation": val_metrics,
        "best_params": best_params,
        "tuning_method": request.tuning_method,
        "cv_folds": request.cv_folds,
        "coefficients": coefficients,
        "warnings": warnings,
        "youden_cutoff": test_metrics.get("youden_cutoff", 0.5),
    }


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok", "version": app.version}


@app.get("/api/demo-data")
def demo_data() -> dict[str, Any]:
    return {"file_name": "Демонстрационный набор данных", **dataset_profile(demo_frame())}


@app.post("/api/dataset/import")
async def import_dataset(file: UploadFile = File(...)) -> dict[str, Any]:
    suffix = Path(file.filename or "").suffix.lower()
    content = await file.read()
    try:
        if suffix == ".csv":
            try:
                df = pd.read_csv(io.BytesIO(content), sep=None, engine="python")
            except UnicodeDecodeError:
                df = pd.read_csv(io.BytesIO(content), encoding="cp1251", sep=None, engine="python")
        elif suffix in {".xlsx", ".xlsm"}:
            df = pd.read_excel(io.BytesIO(content))
        else:
            raise HTTPException(415, "Поддерживаются CSV и XLSX")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, f"Не удалось прочитать файл: {exc}") from exc
    if len(df) > 100_000 or len(df.columns) > 500:
        raise HTTPException(413, "Для MVP лимит: 100 000 строк и 500 столбцов")
    return {"file_name": file.filename, **dataset_profile(df)}


@app.post("/api/analyze/table-one")
def table_one(request: TableOneRequest) -> dict[str, Any]:
    return run_table_one(request)


@app.post("/api/analyze/regression")
def regression(request: RegressionRequest) -> dict[str, Any]:
    return run_regression(request)


class CorrelationRequest(BaseModel):
    rows: list[dict[str, Any]]
    variables: list[str] = Field(min_length=2)
    variable_overrides: dict[str, dict[str, Any]] = {}
    method: str = "auto"  # "pearson" | "spearman" | "auto"


@app.post("/api/analyze/correlation")
def analyze_correlation(request: CorrelationRequest) -> dict[str, Any]:
    from scipy.stats import pearsonr, spearmanr, shapiro
    df = to_frame(request.rows)
    variables = [v for v in request.variables if v in df.columns]
    if len(variables) < 2:
        raise HTTPException(422, "Нужно выбрать хотя бы 2 числовые переменные")

    method = request.method
    if method == "auto":
        try:
            all_normal = all(
                len(df[v].dropna()) >= 3 and shapiro(df[v].dropna().astype(float)).pvalue > 0.05
                for v in variables
            )
            method = "pearson" if all_normal else "spearman"
        except Exception:
            method = "spearman"

    labels = {v: request.variable_overrides.get(v, {}).get("label", v) for v in variables}
    matrix: dict[str, dict] = {}
    for v1 in variables:
        matrix[v1] = {}
        for v2 in variables:
            if v1 == v2:
                matrix[v1][v2] = {"r": 1.0, "p": 0.0, "n": int(df[v1].notna().sum()), "stars": ""}
                continue
            mask = df[v1].notna() & df[v2].notna()
            x = df.loc[mask, v1].astype(float)
            y = df.loc[mask, v2].astype(float)
            n = int(len(x))
            if n < 3:
                matrix[v1][v2] = {"r": None, "p": None, "n": n, "stars": "—"}
                continue
            try:
                r, p = pearsonr(x, y) if method == "pearson" else spearmanr(x, y)
                stars = "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else ""
                matrix[v1][v2] = {"r": round(float(r), 3), "p": round(float(p), 4), "n": n, "stars": stars}
            except Exception:
                matrix[v1][v2] = {"r": None, "p": None, "n": n, "stars": "—"}

    return {
        "method": method,
        "variables": variables,
        "labels": labels,
        "matrix": matrix,
        "n": int(len(df)),
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


@app.post("/api/project/save")
def save_project(request: ProjectSaveRequest) -> dict[str, Any]:
    if request.project_id:
        # overwrite — validate id is safe
        if not re.fullmatch(r"[a-zA-Zа-яА-Я0-9_-]+", request.project_id):
            raise HTTPException(400, "Некорректный project_id")
        project_id = request.project_id
    else:
        slug = re.sub(r"[^a-zA-Zа-яА-Я0-9_-]+", "-", request.project_name).strip("-").lower() or "project"
        project_id = f"{slug[:50]}-{uuid.uuid4().hex[:8]}"
    payload = request.model_dump(exclude={"project_id"})
    payload.update({"project_id": project_id, "saved_at": datetime.now(timezone.utc).isoformat()})
    (PROJECTS_DIR / f"{project_id}.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"project_id": project_id, "saved_at": payload["saved_at"]}


@app.get("/api/projects")
def list_projects() -> list[dict[str, Any]]:
    projects = []
    for path in sorted(PROJECTS_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
            projects.append({key: payload.get(key) for key in ("project_id", "project_name", "file_name", "saved_at")})
        except (OSError, json.JSONDecodeError):
            continue
    return projects


@app.post("/api/project/load")
def load_project(payload: dict[str, str]) -> dict[str, Any]:
    project_id = payload.get("project_id", "")
    if not re.fullmatch(r"[a-zA-Zа-яА-Я0-9_-]+", project_id):
        raise HTTPException(400, "Некорректный идентификатор проекта")
    path = PROJECTS_DIR / f"{project_id}.json"
    if not path.exists():
        raise HTTPException(404, "Проект не найден")
    return json.loads(path.read_text(encoding="utf-8"))


@app.delete("/api/project/{project_id}")
def delete_project(project_id: str) -> dict[str, str]:
    if not re.fullmatch(r"[a-zA-Zа-яА-Я0-9_-]+", project_id):
        raise HTTPException(400, "Некорректный идентификатор проекта")
    path = PROJECTS_DIR / f"{project_id}.json"
    path.unlink(missing_ok=True)  # idempotent: already gone = still OK
    return {"status": "deleted"}


class ScatterRequest(BaseModel):
    x_values: list[float | None]
    y_values: list[float | None]
    x_label: str = ""
    y_label: str = ""
    r: float | None = None
    stars: str = ""
    method: str = "pearson"


def _generate_scatter_bytes(
    x_values: list[float | None],
    y_values: list[float | None],
    x_label: str,
    y_label: str,
    r: float | None,
    stars: str,
    method: str,
    dpi: int = 324,
) -> bytes:
    import warnings
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns

    pairs = [
        (x, y)
        for x, y in zip(x_values, y_values)
        if x is not None and y is not None
        and not (isinstance(x, float) and x != x)
        and not (isinstance(y, float) and y != y)
    ]
    if len(pairs) < 3:
        raise ValueError("Недостаточно данных для графика")

    xs = [p[0] for p in pairs]
    ys = [p[1] for p in pairs]

    sns.set_theme(style="whitegrid", font_scale=0.9)
    fig, ax = plt.subplots(figsize=(4.2, 3.8), dpi=dpi)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        sns.regplot(
            x=xs, y=ys, ax=ax,
            scatter_kws={"alpha": 0.55, "s": 22, "color": "#2563eb", "linewidths": 0.3, "edgecolors": "#1d4ed8"},
            line_kws={"color": "#ef4444", "linewidth": 1.8, "alpha": 0.85},
            ci=95,
        )

    method_label = "Пирсон" if method == "pearson" else "Спирмен"
    title_parts = [method_label]
    if r is not None:
        title_parts.append(f"r = {r:.3f}{stars or ''}")
    ax.set_title("  ".join(title_parts), fontsize=9, pad=6, color="#344054")
    ax.set_xlabel(x_label or "X", fontsize=8, color="#667085")
    ax.set_ylabel(y_label or "Y", fontsize=8, color="#667085")
    ax.tick_params(labelsize=7)
    ax.grid(True, alpha=0.35, linewidth=0.5)
    fig.tight_layout(pad=1.0)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


@app.post("/api/plot/scatter")
def scatter_plot(request: ScatterRequest) -> Response:
    try:
        data = _generate_scatter_bytes(
            request.x_values, request.y_values,
            request.x_label, request.y_label,
            request.r, request.stars or "", request.method,
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    return Response(content=data, media_type="image/png")


def _append_correlation_block(document: Document, corr: CorrelationReportSection) -> None:
    result = corr.result
    variables: list[str] = result["variables"]
    labels: dict[str, str] = result.get("labels", {})
    matrix: dict[str, dict[str, Any]] = result.get("matrix", {})
    method: str = result.get("method", "spearman")
    n: int = result.get("n", 0)
    method_label = "Пирсон (r)" if method == "pearson" else "Спирмен (ρ)"
    symbol = "r" if method == "pearson" else "ρ"

    h = document.add_paragraph()
    run = h.add_run("Корреляционный анализ")
    _apply_tnr(run, 14, bold=True)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(4)

    meta = document.add_paragraph()
    _apply_tnr(meta.add_run(f"Метод: {method_label}. Включено наблюдений: {n}."), 11)
    meta.paragraph_format.space_after = Pt(10)

    if corr.include_matrix and variables:
        n_vars = len(variables)
        tbl = document.add_table(rows=n_vars + 1, cols=n_vars + 1)
        tbl.style = "Table Grid"

        hdr_cells = tbl.rows[0].cells
        _apply_tnr(hdr_cells[0].paragraphs[0].add_run(""), 9)
        for j, var in enumerate(variables):
            p = hdr_cells[j + 1].paragraphs[0]
            _apply_tnr(p.add_run(labels.get(var, var)), 9, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, row_var in enumerate(variables):
            row_cells = tbl.rows[i + 1].cells
            lp = row_cells[0].paragraphs[0]
            _apply_tnr(lp.add_run(labels.get(row_var, row_var)), 9, bold=True)
            for j, col_var in enumerate(variables):
                p = row_cells[j + 1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if row_var == col_var:
                    _apply_tnr(p.add_run("—"), 9)
                else:
                    cell_data = matrix.get(row_var, {}).get(col_var, {})
                    r_val = cell_data.get("r")
                    stars = cell_data.get("stars", "")
                    text = f"{r_val:.3f}{stars}" if r_val is not None else "—"
                    _apply_tnr(p.add_run(text), 9)

        foot = document.add_paragraph()
        _apply_tnr(foot.add_run(f"* p<0,05  ** p<0,01  *** p<0,001  ({symbol} — коэффициент корреляции)"), 9)
        foot.paragraph_format.space_before = Pt(4)
        foot.paragraph_format.space_after = Pt(12)

    if corr.pairs:
        sh = document.add_paragraph()
        _apply_tnr(sh.add_run("Диаграммы рассеяния"), 12, bold=True)
        sh.paragraph_format.space_before = Pt(10)
        sh.paragraph_format.space_after = Pt(6)

        for pair in corr.pairs:
            cell_data = matrix.get(pair.row, {}).get(pair.col, {})
            r_val = cell_data.get("r")
            stars = cell_data.get("stars", "")
            try:
                png = _generate_scatter_bytes(
                    pair.x_values, pair.y_values,
                    pair.x_label or pair.col, pair.y_label or pair.row,
                    r_val, stars, method, dpi=220,
                )
            except (ValueError, Exception):
                continue

            ph = document.add_paragraph()
            ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ph.paragraph_format.space_before = Pt(6)
            ph.paragraph_format.space_after = Pt(2)
            _apply_tnr(ph.add_run(f"{pair.x_label} vs {pair.y_label}"), 11, bold=True)

            img_p = document.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            img_p.paragraph_format.space_after = Pt(10)
            img_p.add_run().add_picture(io.BytesIO(png), width=Inches(3.5))


def _apply_tnr(run: Any, size: float, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), "Times New Roman")


def _normalize_caption(value: str) -> str:
    value = value.strip()
    if not value:
        return "Таблица 1 - Без названия"
    if re.match(r"^Таблица\s+\d+\s*[-–—]", value, flags=re.IGNORECASE):
        return re.sub(r"^(Таблица\s+\d+)\s*[-–—]\s*", r"\1 - ", value, flags=re.IGNORECASE)
    return f"Таблица 1 - {value}"


def _setup_document(document: Document) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)
    rpr = normal_style._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), "Times New Roman")


def _append_table_block(document: Document, request: ExportRequest) -> None:
    analysis = request.analysis
    groups = analysis.get("groups", [])
    caption = document.add_paragraph()
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True
    _apply_tnr(caption.add_run(_normalize_caption(request.title)), 14, False)
    if request.description:
        desc_p = document.add_paragraph()
        desc_p.paragraph_format.space_after = Pt(5)
        _apply_tnr(desc_p.add_run(request.description), 11, False)
    headers = ["Показатель"]
    if request.show_overall:
        headers.append(f"Все (n={analysis.get('n', '—')})")
    headers.extend(f"{g['name']} (n={g['n']})" for g in groups)
    if request.show_missing:
        headers.append("Пропуски")
    if request.show_ci:
        headers.append("95% ДИ")
    headers.append("p-value")
    if request.show_effect:
        headers.append("Размер эффекта")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for item in analysis.get("rows", []):
        has_levels = bool(item.get("levels"))
        cells = table.add_row().cells
        _PRES_SHORT = {"Среднее ± SD": "M ± SD", "Медиана [Q1; Q3]": "Me [Q1; Q3]"}
        row_label = item.get("label") or item.get("variable", "")
        pres = _PRES_SHORT.get(item.get("presentation", ""), item.get("presentation", ""))
        values = [f"{row_label}, {pres}" if pres else row_label]
        if request.show_overall:
            values.append("" if has_levels else item["overall"])
        values.extend("" if has_levels else item.get("groups", {}).get(g["name"], "—") for g in groups)
        if request.show_missing:
            values.append(str(item.get("missing", 0)))
        if request.show_ci:
            values.append(item.get("ci_display", "—"))
        values.append(item.get("p_display", "—"))
        if request.show_effect:
            values.append(f"{item.get('effect_label', '')}: {item.get('effect', '—')}")
        for cell, value in zip(cells, values):
            cell.text = str(value)
        if has_levels:
            for level in item.get("levels", []):
                level_cells = table.add_row().cells
                level_values = [f"   {level.get('level', '—')}"]
                if request.show_overall:
                    level_values.append(level.get("overall", "—"))
                level_values.extend(level.get("groups", {}).get(g["name"], "—") for g in groups)
                if request.show_missing:
                    level_values.append("")
                if request.show_ci:
                    level_values.append("")
                level_values.append("")
                if request.show_effect:
                    level_values.append("")
                for cell, value in zip(level_cells, level_values):
                    cell.text = str(value)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 and row_index > 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _apply_tnr(run, 11 if row_index > 0 else 12, row_index == 0)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(0)
    _apply_tnr(note.add_run(analysis.get("note", "")), 10, False)
    if request.footnotes:
        footnotes_p = document.add_paragraph()
        footnotes_p.paragraph_format.space_before = Pt(3)
        _apply_tnr(footnotes_p.add_run(request.footnotes), 10, False)


def _stream_document(document: Document, filename: str) -> StreamingResponse:
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _append_regression_block(document: Document, payload: dict[str, Any]) -> None:
    analysis = payload.get("analysis") or {}
    labels = payload.get("labels") or {}
    cutoff = float(payload.get("cutoff", 0.5))
    logistic = analysis.get("model_type") == "logistic"
    title = document.add_paragraph()
    _apply_tnr(title.add_run("Регрессионный анализ"), 14, True)
    subtitle = document.add_paragraph()
    outcome = labels.get(analysis.get("outcome"), analysis.get("outcome", "—"))
    model_name = "Бинарная логистическая регрессия" if logistic else "Линейная регрессия"
    _apply_tnr(subtitle.add_run(f"{model_name}. Исход: {outcome}. N = {analysis.get('n', 0)}."), 11, False)
    for warning in analysis.get("warnings") or []:
        warning_p = document.add_paragraph()
        _apply_tnr(warning_p.add_run("Предупреждение. "), 10, True)
        _apply_tnr(warning_p.add_run(str(warning)), 10, False)

    coefficients = analysis.get("coefficients") or []
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Параметр", "β", "SE", "OR" if logistic else "Эффект", "ДИ", "p-value"]
    for cell, text_value in zip(table.rows[0].cells, headers):
        _apply_tnr(cell.paragraphs[0].add_run(text_value), 10, True)
    for row in coefficients:
        cells = table.add_row().cells
        term = str(row.get("term", "—"))
        name, separator, detail = term.partition(": ")
        term_label = f"{labels.get(name, name)}: {detail}" if separator else labels.get(name, name)
        values = [
            term_label,
            fmt_number(float(row.get("estimate", math.nan)), 3),
            fmt_number(float(row.get("standard_error", math.nan)), 3),
            fmt_number(float(row.get("effect", math.nan)), 3),
            f"{fmt_number(float(row.get('effect_ci_lower', math.nan)), 3)}–{fmt_number(float(row.get('effect_ci_upper', math.nan)), 3)}",
            str(row.get("p_display", "—")),
        ]
        for cell, text_value in zip(cells, values):
            _apply_tnr(cell.paragraphs[0].add_run(text_value), 10, False)

    diagnostics = analysis.get("diagnostics") if logistic else None
    if diagnostics:
        image = Image.new("RGB", (800, 480), "white")
        draw = ImageDraw.Draw(image)
        left, top, right, bottom = 82, 38, 752, 402
        draw.line((left, bottom, right, bottom), fill="#52627a", width=3)
        draw.line((left, bottom, left, top), fill="#52627a", width=3)
        draw.line((left, bottom, right, top), fill="#c4ccd8", width=2)
        points = [
            (left + float(point.get("fpr", 0)) * (right - left), bottom - float(point.get("tpr", 0)) * (bottom - top))
            for point in diagnostics.get("roc_curve", [])
        ]
        if len(points) >= 2:
            draw.line(points, fill="#1761c5", width=6, joint="curve")
        draw.text((350, 438), "1 - Specificity", fill="#344054")
        draw.text((12, 208), "Sensitivity", fill="#344054")
        draw.text((570, 52), f"AUC = {float(diagnostics.get('auc', 0)):.3f}", fill="#071d47")
        image_stream = io.BytesIO()
        image.save(image_stream, format="PNG")
        image_stream.seek(0)
        chart_p = document.add_paragraph()
        chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_p.add_run().add_picture(image_stream, width=Mm(130))
        predictions = diagnostics.get("predictions") or []
        tp = tn = fp = fn = 0
        for prediction in predictions:
            positive = float(prediction.get("probability", 0)) >= cutoff
            actual = int(prediction.get("actual", 0))
            if actual == 1 and positive: tp += 1
            elif actual == 0 and not positive: tn += 1
            elif actual == 0 and positive: fp += 1
            else: fn += 1
        ratio = lambda numerator, denominator: numerator / denominator if denominator else 0
        diagnostic_p = document.add_paragraph()
        diagnostic_text = (
            f"ROC AUC = {fmt_number(float(diagnostics.get('auc', math.nan)), 3)}; cut-off = {fmt_number(cutoff, 2)}; "
            f"чувствительность = {fmt_number(ratio(tp, tp + fn) * 100, 1)}%; "
            f"специфичность = {fmt_number(ratio(tn, tn + fp) * 100, 1)}%; "
            f"диагностическая эффективность = {fmt_number(ratio(tp + tn, len(predictions)) * 100, 1)}%."
        )
        _apply_tnr(diagnostic_p.add_run(diagnostic_text), 10, False)
        matrix = document.add_table(rows=3, cols=3)
        matrix.style = "Table Grid"
        matrix_values = [["Факт / прогноз", "+", "−"], ["+", str(tp), str(fn)], ["−", str(fp), str(tn)]]
        for row_index, values in enumerate(matrix_values):
            for cell, value in zip(matrix.rows[row_index].cells, values):
                _apply_tnr(cell.paragraphs[0].add_run(value), 10, row_index == 0 or value in {"+", "−"})


@app.post("/api/dataset/profile")
def compute_profile(request: DatasetProfileRequest) -> dict[str, Any]:
    df = to_frame(request.rows)
    return {"file_name": request.file_name, **dataset_profile(df)}


@app.post("/api/export/table-one.docx")
def export_docx(request: ExportRequest) -> StreamingResponse:
    document = Document()
    _setup_document(document)
    _append_table_block(document, request)
    return _stream_document(document, "table-one.docx")


@app.post("/api/export/report.docx")
def export_report(request: ReportExportRequest) -> StreamingResponse:
    if not request.tables and not request.regression and not request.correlation:
        raise HTTPException(422, "Отчёт не содержит результатов")
    document = Document()
    _setup_document(document)
    for index, table_request in enumerate(request.tables):
        if index > 0:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _append_table_block(document, table_request)
    if request.regression:
        if request.tables:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _append_regression_block(document, request.regression)
    if request.correlation:
        if request.tables or request.regression:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _append_correlation_block(document, request.correlation)
    return _stream_document(document, "report.docx")
